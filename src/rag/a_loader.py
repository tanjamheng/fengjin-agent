"""文档加载器

支持加载多种格式文档：PDF、Markdown、TXT、DOCX
支持递归加载子目录，自动从目录结构提取分类元数据。
"""

from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel
from ..utils.logger import get_logger, generate_trace_id


class Document(BaseModel):
    """文档模型"""
    content: str
    metadata: dict = {}
    source: str = ""


class DocumentLoader:
    """文档加载器"""

    SUPPORTED_EXTENSIONS = {
        ".pdf": "_load_pdf",
        ".md": "_load_markdown",
        ".txt": "_load_text",
        ".docx": "_load_docx"
    }

    def __init__(self, supported_formats: List[str] = None, max_file_size_mb: int = 50):
        self.supported_formats = supported_formats or ["pdf", "md", "txt", "docx"]
        self.max_file_size_mb = max_file_size_mb
        self.log = get_logger(generate_trace_id())

    def load(self, file_path: str, category: str = "") -> Document:
        """加载单个文档

        Args:
            file_path: 文件路径
            category: 知识分类（如 "角色设定"、"世界观"），用于元数据过滤
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValueError(f"文件过大: {file_size_mb:.2f}MB > {self.max_file_size_mb}MB")

        ext = path.suffix.lower()
        method_name = self.SUPPORTED_EXTENSIONS.get(ext)

        if method_name is None:
            raise ValueError(f"不支持的文件格式: {ext}")

        self.log.info(f"加载文档: {file_path}")

        method = getattr(self, method_name)
        content = method(path)

        metadata = {
            "file_name": path.name,
            "file_size_mb": round(file_size_mb, 2),
            "file_ext": ext,
        }
        if category:
            metadata["category"] = category

        return Document(
            content=content,
            metadata=metadata,
            source=str(path)
        )

    def load_directory(self, dir_path: str) -> List[Document]:
        """加载目录下所有文档（仅当前层级，不递归）"""
        path = Path(dir_path)
        if not path.is_dir():
            raise ValueError(f"不是目录: {dir_path}")

        documents = []
        for file_path in sorted(path.iterdir()):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    self.log.warning(f"加载文件失败: {file_path}, 错误: {e}")

        self.log.info(f"从目录 {dir_path} 加载了 {len(documents)} 个文档")
        return documents

    def load_directory_recursive(self, dir_path: str) -> List[Document]:
        """递归加载目录下所有文档，自动从子目录名提取分类元数据

        目录结构约定：
            data/knowledge/
                角色设定/   → category="角色设定"
                世界观/     → category="世界观"
                台词/       → category="台词"
                人物关系/   → category="人物关系"
                剧情事件/   → category="剧情事件"

        Args:
            dir_path: 知识库根目录路径
        """
        root = Path(dir_path)
        if not root.is_dir():
            raise ValueError(f"不是目录: {dir_path}")

        documents = []

        # 遍历子目录
        for subdir in sorted(root.iterdir()):
            if not subdir.is_dir():
                continue

            category = subdir.name
            for file_path in sorted(subdir.iterdir()):
                if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    try:
                        doc = self.load(str(file_path), category=category)
                        documents.append(doc)
                    except Exception as e:
                        self.log.warning(f"加载文件失败: {file_path}, 错误: {e}")

        # 也检查根目录下的文件（无分类）
        for file_path in sorted(root.iterdir()):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    self.log.warning(f"加载文件失败: {file_path}, 错误: {e}")

        self.log.info(f"递归加载 {dir_path}，共 {len(documents)} 个文档")
        return documents

    def _load_text(self, path: Path) -> str:
        """加载纯文本"""
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    def _load_markdown(self, path: Path) -> str:
        """加载 Markdown"""
        return self._load_text(path)

    def _load_pdf(self, path: Path) -> str:
        """加载 PDF"""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            self.log.warning("pypdf 未安装，尝试使用 PyMuPDF")
            try:
                import fitz
                doc = fitz.open(str(path))
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                return text
            except ImportError:
                raise ImportError("请安装 pypdf 或 PyMuPDF: pip install pypdf 或 pip install PyMuPDF")

    def _load_docx(self, path: Path) -> str:
        """加载 DOCX"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
